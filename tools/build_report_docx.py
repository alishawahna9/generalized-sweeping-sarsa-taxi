from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reports" / "final_report.md"
OUTPUT = ROOT / "reports" / "final_report.docx"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(18)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Courier New"
    code.font.size = Pt(9)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.left_indent = Inches(0.15)


def add_text_with_inline_code(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            add_basic_markdown_runs(paragraph, part)


def add_basic_markdown_runs(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="Code Block")
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def add_simple_table(doc, rows):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.autofit = False
    set_table_borders(table)

    width = Inches(6.5)
    col_width = width / col_count
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx in range(col_count):
            text = row[col_idx].strip() if col_idx < len(row) else ""
            cells[col_idx].width = col_width
            set_cell_margins(cells[col_idx])
            paragraph = cells[col_idx].paragraphs[0]
            add_text_with_inline_code(paragraph, text)
            for run in paragraph.runs:
                run.font.size = Pt(9.5)
                if row_idx == 0:
                    run.bold = True
    doc.add_paragraph()


def parse_table(lines):
    rows = []
    for line in lines:
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        if all(set(part) <= {"-", ":", " "} for part in parts):
            continue
        rows.append(parts)
    return rows


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    markdown = SOURCE.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        raw = lines[i].rstrip()

        if raw.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if raw.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_simple_table(doc, parse_table(table_lines))
            continue

        if not raw.strip():
            i += 1
            continue

        if raw.startswith("# "):
            p = doc.add_paragraph(style="Title")
            add_text_with_inline_code(p, raw[2:])
        elif raw.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_text_with_inline_code(p, raw[3:])
        elif raw.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_text_with_inline_code(p, raw[4:])
        elif raw.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_text_with_inline_code(p, raw[2:])
        else:
            p = doc.add_paragraph()
            add_text_with_inline_code(p, raw)

        i += 1

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
